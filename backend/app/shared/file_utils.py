"""
Shared file utilities
Low-level utility functions for PDF, DOCX file processing
"""

import logging
from pathlib import Path
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)


def pdf_to_text(file_path: str) -> Dict[str, Any]:
    """
    Extract text from PDF file

    Args:
        file_path: PDF file path

    Returns:
        Dict containing:
            - text: Extracted text content
            - metadata: Information including page_count, etc.
    """
    try:
        import PyPDF2

        text_parts = []
        page_count = 0

        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            page_count = len(pdf_reader.pages)

            for page_num in range(page_count):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                text_parts.append(text)

        full_text = "\n\n".join(text_parts)
        word_count = len(full_text.split())

        return {
            "text": full_text,
            "metadata": {
                "page_count": page_count,
                "word_count": word_count,
                "file_path": file_path
            }
        }

    except ImportError:
        logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
        raise Exception("PDF processing library not available")
    except Exception as e:
        logger.error(f"Failed to extract text from PDF {file_path}: {e}")
        raise


def docx_to_text(file_path: str) -> Dict[str, Any]:
    """
    Extract text from DOCX file

    Args:
        file_path: DOCX file path

    Returns:
        Dict containing:
            - text: Extracted text content
            - metadata: Information including paragraph_count, etc.
    """
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []
        paragraph_count = 0

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
                paragraph_count += 1

        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = "\n\n".join(text_parts)
        word_count = len(full_text.split())

        return {
            "text": full_text,
            "metadata": {
                "paragraph_count": paragraph_count,
                "word_count": word_count,
                "file_path": file_path
            }
        }

    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        raise Exception("DOCX processing library not available")
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
        raise


def detect_file_type(file_path: str) -> str:
    """
    Detect file type

    Args:
        file_path: File path

    Returns:
        File type: 'pdf', 'docx', 'doc', or others
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == '.pdf':
        return 'pdf'
    elif suffix in ['.docx', '.doc']:
        return 'docx' if suffix == '.docx' else 'doc'
    else:
        return suffix.lstrip('.')


def get_file_size(file_path: str) -> int:
    """
    Get file size in bytes

    Args:
        file_path: File path

    Returns:
        File size in bytes
    """
    return Path(file_path).stat().st_size


def pdf_has_text_layer(file_path: str) -> bool:
    """
    Check if PDF has extractable text layer
    Scanned PDFs typically have no text layer

    Args:
        file_path: PDF file path

    Returns:
        True if PDF has text layer, False if scanned/image-based
    """
    try:
        import PyPDF2

        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            # Check first few pages for text content
            pages_to_check = min(3, len(pdf_reader.pages))
            total_text_length = 0

            for page_num in range(pages_to_check):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                total_text_length += len(text.strip())

            # If we found substantial text, assume it has text layer
            # Threshold: at least 50 characters across first 3 pages
            return total_text_length >= 50

    except Exception as e:
        logger.warning(f"Failed to check PDF text layer for {file_path}: {e}")
        # On error, assume it might be scanned (safer to try OCR)
        return False
